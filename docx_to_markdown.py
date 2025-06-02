from docx import Document
import os
import re

def docx_to_markdown(docx_path):
    doc = Document(docx_path)
    markdown_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Check for heading styles
        if para.style.name.startswith('Heading'):
            level = int(re.search(r'\d', para.style.name).group()) if re.search(r'\d', para.style.name) else 1
            markdown_content.append('#' * level + ' ' + text)
        elif para.style.name == 'Title':
            markdown_content.append('# ' + text)
        else:
            # Check for list items
            if para.style.name.startswith('List'):
                markdown_content.append('- ' + text)
            else:
                # Handle bold and italic
                if para.runs:
                    formatted_text = ""
                    for run in para.runs:
                        run_text = run.text
                        if run.bold and run.italic:
                            formatted_text += f"***{run_text}***"
                        elif run.bold:
                            formatted_text += f"**{run_text}**"
                        elif run.italic:
                            formatted_text += f"*{run_text}*"
                        else:
                            formatted_text += run_text
                    markdown_content.append(formatted_text)
                else:
                    markdown_content.append(text)
        
        markdown_content.append('')  # Add blank line between paragraphs
    
    # Handle tables
    for table in doc.tables:
        markdown_content.append('')
        
        # Header row
        if table.rows:
            header_cells = [cell.text.strip() for cell in table.rows[0].cells]
            markdown_content.append('| ' + ' | '.join(header_cells) + ' |')
            markdown_content.append('|' + '---|' * len(header_cells))
            
            # Data rows
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                markdown_content.append('| ' + ' | '.join(cells) + ' |')
        
        markdown_content.append('')
    
    return '\n'.join(markdown_content)

# Convert all docx files
reports_dir = '/Users/deepwork/LSG AI Audit/Reports'
output_dir = '/Users/deepwork/LSG AI Audit/markdown_reports'

os.makedirs(output_dir, exist_ok=True)

docx_files = [f for f in os.listdir(reports_dir) if f.endswith('.docx')]

for docx_file in docx_files:
    docx_path = os.path.join(reports_dir, docx_file)
    markdown_content = docx_to_markdown(docx_path)
    
    # Create markdown filename
    md_filename = docx_file.replace('.docx', '.md')
    md_path = os.path.join(output_dir, md_filename)
    
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(markdown_content)
    
    print(f"Converted: {docx_file} -> {md_filename}")